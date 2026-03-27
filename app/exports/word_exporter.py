from io import BytesIO

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.exports.institutional_layout import INSTITUTIONAL_GREEN, TABLE_HEADERS, get_logo_path
from app.exports.report_builder import ReportBuilder


class WordExporter:
    @staticmethod
    def export_group_schedule(group_id: int) -> BytesIO:
        context = ReportBuilder.build_group_report_context(group_id)
        document = Document()
        WordExporter._configure_document(document)
        WordExporter._render_header(document, context)
        WordExporter._render_table(document, context)

        output = BytesIO()
        document.save(output)
        output.seek(0)
        return output

    @staticmethod
    def _configure_document(document: Document) -> None:
        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)

    @staticmethod
    def _render_header(document: Document, context: dict) -> None:
        # Header table with logo, institution info, and date/page
        table = document.add_table(rows=4, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        # Column widths: time | logo | institution info (expanded) | | date/page
        widths = [Cm(3.2), Cm(1.5), Cm(13.0), Cm(3.5), Cm(4.0)]
        for column_index, width in enumerate(widths):
            for cell in table.columns[column_index].cells:
                cell.width = width

        # Left column: time and report code
        table.cell(0, 0).text = context["generated_time"]
        table.cell(1, 0).text = context["report_code"]
        _format_paragraph(table.cell(0, 0).paragraphs[0], size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
        _format_paragraph(table.cell(1, 0).paragraphs[0], size=8, align=WD_ALIGN_PARAGRAPH.LEFT)

        # Logo column (merge rows 0-3)
        logo_cell = table.cell(0, 1).merge(table.cell(3, 1))
        logo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        logo_paragraph = logo_cell.paragraphs[0]
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        logo_path = get_logo_path()
        if logo_path is not None:
            logo_paragraph.add_run().add_picture(str(logo_path), height=Cm(2.5))
        else:
            logo_run = logo_paragraph.add_run("LOGO")
            logo_run.bold = True
            logo_run.font.size = Pt(10)

        # Institution info (merge columns 2-3, rows 0-3)
        info_cell = table.cell(0, 2).merge(table.cell(3, 3))
        info_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # Clear default paragraph and add custom ones
        for _ in info_cell.paragraphs:
            p = _
            for run in p.runs:
                r = run._r
                r.getparent().remove(r)
        
        # Institution name
        p1 = info_cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(context["institution_name"])
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.name = "Arial"
        r1.font.color.rgb = RGBColor.from_string(INSTITUTIONAL_GREEN)
        
        # Coordination
        p2 = info_cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(context["coordination_name"])
        r2.font.size = Pt(9)
        r2.font.name = "Arial"
        r2.font.color.rgb = RGBColor.from_string("666666")
        
        # Report title
        p3 = info_cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(context["report_title"])
        r3.font.size = Pt(9)
        r3.font.name = "Arial"
        r3.font.color.rgb = RGBColor.from_string("666666")
        
        # Period
        p4 = info_cell.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r4 = p4.add_run(f"Periodo: {context['period']}")
        r4.font.size = Pt(8)
        r4.font.name = "Arial"
        r4.font.color.rgb = RGBColor.from_string("666666")

        # Right column: date and page
        table.cell(0, 4).text = context["generated_date"]
        table.cell(1, 4).text = "PAG. 1"
        _format_paragraph(table.cell(0, 4).paragraphs[0], size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _format_paragraph(table.cell(1, 4).paragraphs[0], size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)

        # Remove borders from header table
        _remove_table_borders(table)

        # Meta information table
        meta = document.add_table(rows=3, cols=6)
        meta.alignment = WD_TABLE_ALIGNMENT.LEFT
        meta.autofit = False
        meta_values = [
            ["UNIDAD ACADÉMICA:", context["unit_code"], context["unit_name"], "", "GRUPO:", str(context["group_number"])],
            ["CARRERA:", context["career_code"], context["career_name"], "", "", ""],
            ["PLAN DE ESTUDIOS:", context["plan_key"], "", "", "", ""],
        ]
        widths = [Cm(3.8), Cm(1.2), Cm(8.5), Cm(2.0), Cm(2.5), Cm(2.0)]
        for row_index, row_values in enumerate(meta_values):
            for col_index, value in enumerate(row_values):
                cell = meta.cell(row_index, col_index)
                cell.width = widths[col_index]
                cell.text = value
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                _format_paragraph(cell.paragraphs[0], bold=col_index in {0, 4}, size=10)

        document.add_paragraph()

    @staticmethod
    def _render_table(document: Document, context: dict) -> None:
        table = document.add_table(rows=1, cols=len(TABLE_HEADERS))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        header_cells = table.rows[0].cells
        for index, header in enumerate(TABLE_HEADERS):
            header_cells[index].text = header
            header_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_shading(header_cells[index], INSTITUTIONAL_GREEN)
            _format_paragraph(header_cells[index].paragraphs[0], bold=True, size=7, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)

        for row_data in context["rows"]:
            row_cells = table.add_row().cells
            values = [
                row_data["clave_control"],
                row_data["descripcion_maestro"],
                row_data["edificio"],
                row_data["salon"],
                str(row_data["cap_asg"]),
                str(row_data["tpo_sgp"]),
                row_data["lunes"],
                row_data["martes"],
                row_data["miercoles"],
                row_data["jueves"],
                row_data["viernes"],
                row_data["sabado"],
                row_data["domingo"],
                str(row_data["es"]),
            ]
            for index, value in enumerate(values):
                row_cells[index].text = value
                row_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                align = WD_ALIGN_PARAGRAPH.LEFT if index in {0, 1} else WD_ALIGN_PARAGRAPH.CENTER
                _format_paragraph(row_cells[index].paragraphs[0], size=7, align=align)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _remove_table_borders(table) -> None:
    """Remove all borders from a table"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _format_paragraph(paragraph, bold: bool = False, size: int = 9, color: str = "000000", align=None) -> None:
    if align is not None:
        paragraph.alignment = align
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.color.rgb = RGBColor.from_string(color)
