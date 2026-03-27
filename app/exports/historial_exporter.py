"""Exportador de historial de docentes (Excel, Word, PDF)."""
from datetime import datetime
from io import BytesIO

# openpyxl
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

# python-docx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

# reportlab
from reportlab.lib import colors
from reportlab.lib.pagesizes import landscape, letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.exports.institutional_layout import (
    BLACK,
    INSTITUTIONAL_GRAY,
    INSTITUTIONAL_GREEN,
    INSTITUTIONAL_GREEN_LIGHT,
    TABLE_GRID,
    WHITE,
    INSTITUTION_NAME,
)

_HIST_HEADERS = ["Docente", "Clave Docente", "Materia", "Clave Materia", "Grupos", "Horas", "Observación", "Nivel"]
_NIVEL_LABELS = {"malo": "Da mal la materia", "regular": "Regular", "bueno": "La da bien"}
_NIVEL_COLORS_EXCEL = {"malo": "FFCCCC", "regular": "FFF3CC", "bueno": "CCFFCC"}

_NOW_STR = lambda: datetime.now().strftime("%d/%m/%Y %H:%M")


def _row_data(entry: dict) -> list:
    doc = entry.get("docente") or {}
    mat = entry.get("materia") or {}
    obs = entry.get("observacion") or {}
    nivel = obs.get("nivel", "") if obs else ""
    return [
        doc.get("nombre", ""),
        doc.get("clave_docente", ""),
        mat.get("nombre", ""),
        mat.get("clave", ""),
        ", ".join(entry.get("grupos", [])),
        str(entry.get("horas_asignadas", 0)),
        obs.get("observacion", "") if obs else "",
        _NIVEL_LABELS.get(nivel, nivel),
    ]


# ──────────────────────────────────────────────────────────────────────────
#  Excel
# ──────────────────────────────────────────────────────────────────────────

class HistorialExcelExporter:
    @staticmethod
    def export(rows: list[dict], title: str = "Historial de Docentes") -> BytesIO:
        wb = Workbook()
        ws = wb.active
        ws.title = "Historial"
        ws.sheet_view.showGridLines = False

        green_fill = PatternFill("solid", fgColor=INSTITUTIONAL_GREEN)
        light_fill = PatternFill("solid", fgColor=INSTITUTIONAL_GREEN_LIGHT)
        white_bold = Font(color=WHITE, bold=True, size=11)
        bold_font = Font(bold=True, size=10)
        normal_font = Font(size=10)
        gray_font = Font(color=INSTITUTIONAL_GRAY, size=9)
        thin = Border(
            left=Side(style="thin", color=TABLE_GRID),
            right=Side(style="thin", color=TABLE_GRID),
            top=Side(style="thin", color=TABLE_GRID),
            bottom=Side(style="thin", color=TABLE_GRID),
        )
        center = Alignment(horizontal="center", vertical="center", wrap_text=True)
        left = Alignment(horizontal="left", vertical="center", wrap_text=True)

        # Title row
        ws.merge_cells("A1:H1")
        ws["A1"] = title
        ws["A1"].font = Font(bold=True, size=13, color=INSTITUTIONAL_GREEN)
        ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 24

        # Generated at row
        ws.merge_cells("A2:H2")
        ws["A2"] = f"Generado: {_NOW_STR()}  |  {INSTITUTION_NAME}"
        ws["A2"].font = gray_font
        ws["A2"].alignment = left
        ws.row_dimensions[2].height = 16

        # Header row
        cols = ["A", "B", "C", "D", "E", "F", "G", "H"]
        widths = [32, 14, 32, 14, 14, 8, 40, 20]
        for col, w in zip(cols, widths):
            ws.column_dimensions[col].width = w

        for idx, (col, header) in enumerate(zip(cols, _HIST_HEADERS)):
            cell = ws[f"{col}3"]
            cell.value = header
            cell.font = white_bold
            cell.fill = green_fill
            cell.border = thin
            cell.alignment = center
        ws.row_dimensions[3].height = 20

        # Data rows
        for i, entry in enumerate(rows):
            row_num = i + 4
            data = _row_data(entry)
            obs = entry.get("observacion") or {}
            nivel = obs.get("nivel", "") if obs else ""
            nivel_fill = PatternFill("solid", fgColor=_NIVEL_COLORS_EXCEL.get(nivel, "FFFFFF")) if nivel else None

            for j, (col, value) in enumerate(zip(cols, data)):
                cell = ws[f"{col}{row_num}"]
                cell.value = value
                cell.border = thin
                cell.alignment = left
                cell.font = normal_font
                if nivel_fill and j in (6, 7):
                    cell.fill = nivel_fill
                elif i % 2 == 1:
                    cell.fill = light_fill

            ws.row_dimensions[row_num].height = 18

        output = BytesIO()
        wb.save(output)
        output.seek(0)
        return output


# ──────────────────────────────────────────────────────────────────────────
#  Word
# ──────────────────────────────────────────────────────────────────────────

class HistorialWordExporter:
    @staticmethod
    def export(rows: list[dict], title: str = "Historial de Docentes") -> BytesIO:
        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(35)
        section.page_height = Cm(22)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

        # Title
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x00, 0x6B, 0x3F)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub = doc.add_paragraph(f"Generado: {_NOW_STR()}  |  {INSTITUTION_NAME}")
        sub.runs[0].font.size = Pt(8)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Table
        table = doc.add_table(rows=1 + len(rows), cols=len(_HIST_HEADERS))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        header_row = table.rows[0]
        green_rgb = RGBColor(0x00, 0x6B, 0x3F)
        for idx, header in enumerate(_HIST_HEADERS):
            cell = header_row.cells[idx]
            cell.text = header
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_bg(cell, "006B3F")

        _nivel_colors = {"malo": "FFCCCC", "regular": "FFF3CC", "bueno": "CCFFCC"}
        for i, entry in enumerate(rows):
            data = _row_data(entry)
            obs = entry.get("observacion") or {}
            nivel = obs.get("nivel", "") if obs else ""
            data_row = table.rows[i + 1]
            for j, val in enumerate(data):
                cell = data_row.cells[j]
                cell.text = str(val)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.runs[0] if p.runs else p.add_run(str(val))
                run.font.size = Pt(7)
                if nivel and j in (6, 7):
                    _set_cell_bg(cell, _nivel_colors.get(nivel, "FFFFFF"))
                elif i % 2 == 1:
                    _set_cell_bg(cell, "D9EEE3")

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output


def _set_cell_bg(cell, hex_color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


# ──────────────────────────────────────────────────────────────────────────
#  PDF
# ──────────────────────────────────────────────────────────────────────────

class HistorialPdfExporter:
    @staticmethod
    def export(rows: list[dict], title: str = "Historial de Docentes") -> BytesIO:
        output = BytesIO()
        doc = SimpleDocTemplate(
            output,
            pagesize=landscape(letter),
            leftMargin=1.0 * cm,
            rightMargin=1.0 * cm,
            topMargin=1.5 * cm,
            bottomMargin=1.0 * cm,
        )

        styles = getSampleStyleSheet()
        small = ParagraphStyle("small", parent=styles["BodyText"], fontName="Helvetica", fontSize=7, leading=8)
        small_c = ParagraphStyle("smallc", parent=small, alignment=1)
        header_s = ParagraphStyle("hdr", parent=small, fontName="Helvetica-Bold", textColor=colors.white, alignment=1)

        table_data = [[Paragraph(h, header_s) for h in _HIST_HEADERS]]

        _pdf_nivels = {
            "malo": colors.HexColor("#FFCCCC"),
            "regular": colors.HexColor("#FFF3CC"),
            "bueno": colors.HexColor("#CCFFCC"),
        }

        row_styles = []
        for i, entry in enumerate(rows):
            data = _row_data(entry)
            obs = entry.get("observacion") or {}
            nivel = obs.get("nivel", "") if obs else ""
            table_data.append([
                Paragraph(data[0], small),
                Paragraph(data[1], small_c),
                Paragraph(data[2], small),
                Paragraph(data[3], small_c),
                Paragraph(data[4], small_c),
                Paragraph(data[5], small_c),
                Paragraph(data[6], small),
                Paragraph(data[7], small_c),
            ])
            if nivel:
                bg = _pdf_nivels.get(nivel, colors.white)
                row_styles.append(("BACKGROUND", (6, i + 1), (7, i + 1), bg))
            elif i % 2 == 1:
                row_styles.append(("BACKGROUND", (0, i + 1), (-1, i + 1), colors.HexColor("#D9EEE3")))

        col_widths = [5.5*cm, 2.5*cm, 5.5*cm, 2.5*cm, 2.5*cm, 1.5*cm, 6.5*cm, 3.0*cm]
        tbl = Table(table_data, repeatRows=1, colWidths=col_widths)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{INSTITUTIONAL_GREEN}")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#B7B7B7")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            *row_styles,
        ]))

        title_p = Paragraph(f"<b>{title}</b>", ParagraphStyle(
            "title_p", parent=styles["BodyText"],
            fontName="Helvetica-Bold", fontSize=12,
            textColor=colors.HexColor(f"#{INSTITUTIONAL_GREEN}"),
            alignment=1, spaceAfter=4,
        ))
        sub_p = Paragraph(
            f"Generado: {_NOW_STR()}  |  {INSTITUTION_NAME}",
            ParagraphStyle("sub_p", parent=styles["BodyText"], fontName="Helvetica", fontSize=8,
                           textColor=colors.HexColor(f"#{INSTITUTIONAL_GRAY}"), alignment=1, spaceAfter=8),
        )

        story = [title_p, sub_p, tbl]
        doc.build(story)
        output.seek(0)
        return output
